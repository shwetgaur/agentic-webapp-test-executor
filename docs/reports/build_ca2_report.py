"""Build CA-2 Monthly Progress Report (August 2026) as a .docx."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "CA2_Monthly_Progress_Report_August_2026_DS1.docx"
DOWNLOADS = Path.home() / "Downloads" / "CA2_Monthly_Progress_Report_August_2026_DS1.docx"
FIG = Path(__file__).resolve().parent / "fig_architecture.png"


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(OxmlElement("w:instrText"))
    run._r.append(instr)
    run._r.append(fld2)
    set_run_font(run, size=10)


def shade_header_row(row):
    for cell in row.cells:
        tc = cell._tePr if hasattr(cell, "_tePr") else cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F4E79")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
                set_run_font(run, size=10, bold=True)


def set_cell_text(cell, text, bold=False, size=11, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=10, align="center")
    shade_header_row(table.rows[0])
    for r_i, row in enumerate(rows, 1):
        for c, val in enumerate(row):
            set_cell_text(table.rows[r_i].cells[c], str(val), size=10)
            table.rows[r_i].cells[c].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def body(doc, text, first_line_indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(8)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    set_run_font(run)
    return p


def heading(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        set_run_font(run, size=14, bold=True)
        run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def heading2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        set_run_font(run, size=13, bold=True)
        run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=10, italic=True)
    return p


def make_architecture_fig():
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch

    fig, ax = plt.subplots(figsize=(10.2, 5.6))
    ax.set_xlim(0, 10.2)
    ax.set_ylim(0, 5.6)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    boxes = [
        (0.25, 2.2, 1.7, 1.3, "Structured\nTest Prompt\n(YAML / form)", "#D6EAF8"),
        (2.2, 2.2, 1.8, 1.3, "Parser &\nValidator\n(rule-based)", "#D5F5E3"),
        (4.25, 2.2, 1.9, 1.3, "Playwright\nExecutor\n(Chromium)", "#FCF3CF"),
        (6.4, 3.35, 1.8, 1.15, "JSON + MD\nReport\n+ screenshot", "#FAD7A0"),
        (6.4, 1.15, 1.8, 1.15, "Notify Agent\n(team map +\nticket id)", "#F5B7B1"),
        (8.5, 2.2, 1.7, 1.3, "CLI / API /\nStreamlit\nDemo", "#E8DAEF"),
    ]
    for x, y, w, h, label, color in boxes:
        ax.add_patch(
            FancyBboxPatch(
                (x, y), w, h, boxstyle="round,pad=0.03,rounding_size=0.12",
                facecolor=color, edgecolor="#1F4E79", linewidth=1.4,
            )
        )
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=8.5, color="#1F1F1F")

    arrows = [
        ((1.95, 2.85), (2.2, 2.85)),
        ((4.0, 2.85), (4.25, 2.85)),
        ((6.15, 3.15), (6.4, 3.7)),
        ((6.15, 2.55), (6.4, 1.7)),
        ((8.2, 3.9), (8.5, 2.95)),
        ((8.2, 1.7), (8.5, 2.7)),
    ]
    for (x1, y1), (x2, y2) in arrows:
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="->", color="#1F4E79", lw=1.3))

    ax.set_title("Fig. 1  Pipeline of the Agentic Web-App Test Executor (MVP)", fontsize=10, pad=8)
    fig.tight_layout()
    fig.savefig(FIG, dpi=180, bbox_inches="tight")
    plt.close()


def build():
    make_architecture_fig()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = header.add_run(
        "Department of Artificial Intelligence & Machine Learning, SIT Pune  |  B.Tech Project  |  Sem VII  |  AY 2026-27"
    )
    set_run_font(hr, size=9)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("CA-2 Monthly Progress Report  •  August 2026  •  Page ")
    set_run_font(fr, size=10)
    add_page_number(footer)

    # -------- COVER --------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Project ID: DS-1 (Dassault Systèmes)")
    set_run_font(r, size=12, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SYMBIOSIS INSTITUTE OF TECHNOLOGY, PUNE")
    set_run_font(r, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Department of Artificial Intelligence and Machine Learning")
    set_run_font(r, size=12, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("B. Tech. Project")
    set_run_font(r, size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Agentic Web-App Test Executor")
    set_run_font(r, size=20, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("(Domain: Quality Engineering)")
    set_run_font(r, size=13, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Monthly Progress Report – August 2026")
    set_run_font(r, size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CA-2  |  First Monthly Progress Report  |  Semester VII")
    set_run_font(r, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Industry Problem Statement: DS 1  |  Partner: Dassault Systèmes (ENOVIA)")
    set_run_font(r, size=12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Under the Guidance of")
    set_run_font(r, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prof. Mayur Gaikwad")
    set_run_font(r, size=13)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("(Name and sign of the Guide)")
    set_run_font(r, size=11, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Industry Mentor: To be confirmed in Dassault mentor meeting")
    set_run_font(r, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Group Members")
    set_run_font(r, size=14, bold=True)

    add_table(
        doc,
        ["Sr.", "Name of Student", "PRN"],
        [
            ["1", "Shwet Gaur", "________________"],
            ["2", "Sahishnu Raut", "________________"],
            ["3", "Eesha Barad", "________________"],
            ["4", "Saksham Sharma", "________________"],
        ],
        [0.7, 3.5, 2.5],
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DEPARTMENT OF ARTIFICIAL INTELLIGENCE AND MACHINE LEARNING")
    set_run_font(r, size=12, bold=True)

    doc.add_page_break()

    # -------- TOC --------
    heading(doc, "TABLE OF CONTENTS")
    add_table(
        doc,
        ["SN", "Particulars", "Page"],
        [
            ["1", "Names of all the group members and their work distribution", "3"],
            ["2", "Introduction", ""],
            ["3", "Problem Statement", ""],
            ["4", "Objectives of the Project", ""],
            ["5", "Literature Review", ""],
            ["6", "Gap in the Research / Technology / Methodology", ""],
            ["7", "Description of the proposed solution", ""],
            ["8", "Requirement Analysis", ""],
            ["9", "Technology Stack", ""],
            ["10", "Design", ""],
            ["11", "Development / Implementation", ""],
            ["12", "Testing and Debugging", ""],
            ["13", "Project plan and progress", ""],
            ["14", "Project Outcome", ""],
            ["15", "Conclusion", ""],
            ["16", "References (APA format)", ""],
            ["A", "Appendix A – Similarity Report (to be attached by Guide)", ""],
            ["B", "Appendix B – AI Plagiarism Report (to be attached by Guide)", ""],
        ],
    )
    body(
        doc,
        "Note: Update page numbers in Word using References → Table of Contents after the final print. "
        "Fill PRNs on the cover page before Moodle upload. Similarity and AI reports will be generated by the Guide and must be attached at the end.",
        first_line_indent=False,
    )

    # 1
    heading(doc, "1. Names of Group Members and Work Distribution")
    body(
        doc,
        "This CA-2 report covers work done after problem allocation (DS 1 – Agentic Web-App Test Executor) up to the last week of August 2026. "
        "The team has four members. Tasks were split so that each person owns one module, but the full pipeline was also integrated as a single demoable system.",
    )
    add_table(
        doc,
        ["Name", "Role (this month)", "Work done since last evaluation"],
        [
            [
                "Shwet Gaur",
                "Executor + integration",
                "Playwright runner (goto, fill, click, select, assert); failure screenshots; CLI run_suite.py; end-to-end wiring of parser → execute → report → notify; Streamlit demo walkthrough.",
            ],
            [
                "Sahishnu Raut",
                "Parser / AI contract",
                "Step schema and structured prompt fields; rule-based NL parser; YAML sample cases; architecture diagram; literature shortlist for CA-2.",
            ],
            [
                "Eesha Barad",
                "Platform / reports",
                "Pydantic models; FastAPI endpoints (/health, /run/text, /run/structured); JSON and Markdown report writer; report folder layout.",
            ],
            [
                "Saksham Sharma",
                "Notify + QA + docs",
                "Team ownership YAML; notify agent (console + optional Slack); 10 sample text cases; logbook coordination; this monthly report draft.",
            ],
        ],
        [1.6, 1.7, 4.2],
    )
    body(
        doc,
        "All four members reviewed the Dassault brief, met the faculty guide, and are expected to be able to explain the full pipeline in the CA-2 discussion. "
        "Weekly logbook pages are being maintained and shown to the Project Coordinator as instructed.",
    )

    # 2
    heading(doc, "2. Introduction")
    body(
        doc,
        "Web products in industry are tested again and again after every release. A large part of this work is still written as English steps in test cases: open a URL, enter a user name, click Login, check that the next page is shown. "
        "When the same steps are automated, testers usually rewrite them as Selenium or Playwright scripts. Those scripts take time to write and they break when a button id or a label is changed.",
    )
    body(
        doc,
        "Dassault Systèmes has given this as DS 1 under Quality Engineering. The ask is simple to state: take basic text steps, replay them on a live web application, and say whether the application behaved as expected. "
        "The faculty guide also asked that each run should produce a written record of pass and fail, and that a failed case should be sent to the team that owns that feature, similar to how a scrum master would route a defect.",
    )
    body(
        doc,
        "Our project tries to join these pieces in one pipeline instead of using five separate tools. In this reporting period we have frozen the input format, built a first working executor, stored reports on disk, and shown a failure alert with a ticket id. "
        "The parser is still rule-based. An LLM layer is planned next, once the contracts are stable. This report records that Stage-1 work.",
    )

    # 3
    heading(doc, "3. Problem Statement")
    body(
        doc,
        "Design and implement an agentic test executor that accepts basic text-based test steps (and a structured test form), maps them to browser actions, runs the flow on a web application using Playwright, verifies expected UI behaviour, writes a step-level pass/fail report with evidence, and notifies the mapped owning team when a test fails.",
    )
    heading2(doc, "3.1 Industry context")
    body(
        doc,
        "Manual UI checking is slow and two testers may not follow the same path. Recorded scripts are faster but they depend on locators. When the user interface is redesigned, many scripts fail even if the product still works. "
        "Commercial tools such as Testim try to reduce this pain with smart locators. Microsoft has also started shipping Playwright Test Agents that plan, generate and heal tests. Our college project cannot copy a commercial product. It can, however, show a clear academic prototype on a public demo site (Sauce Demo) with contracts that later LLM work can reuse.",
    )

    # 4
    heading(doc, "4. Objectives of the Project")
    body(doc, "The objectives below are frozen for Semester VII. They include the two extra items asked by the guide (documentation and notify).", first_line_indent=False)
    add_table(
        doc,
        ["No.", "Objective", "Status in Aug 2026"],
        [
            ["O1", "Convert basic text / structured steps into a fixed JSON action list.", "Done (rule-based)"],
            ["O2", "Execute those actions on a web app through Playwright (Chromium).", "Done (MVP actions)"],
            ["O3", "Verify UI state using assertions on text, URL and visibility.", "Done"],
            ["O4", "Write a structured run report (JSON + Markdown) with screenshots on fail.", "Done"],
            ["O5", "Route a failure alert to the owning team using a module–team map and a ticket id.", "Done (console; Slack optional)"],
            ["O6", "Reduce script brittleness using AI-assisted parsing / locator recovery.", "Not started (Phase-2/3)"],
            ["O7", "Evaluate on sample apps and keep college + Dassault documents in order.", "In progress"],
        ],
        [0.7, 4.6, 2.2],
    )

    # 5
    heading(doc, "5. Literature Review")
    body(
        doc,
        "Five sources were studied in this period. The aim was not to copy any tool, but to see where DS 1 sits between industrial products and research papers. Each source is summarised in our own words. Full APA entries are listed in Section 16.",
    )

    heading2(doc, "5.1 Playwright Test Agents (Microsoft, 2025–26)")
    body(
        doc,
        "Playwright now describes three agents: a planner that walks an application and writes a Markdown plan, a generator that turns the plan into test files, and a healer that tries to repair a failing test. "
        "This is the closest official stack to DS 1. The difference is that we do not generate a .spec.ts file as the main output. We keep a JSON step list, run it at once, and attach a report plus a notify event. That choice makes the college demo easier to explain in one sitting.",
    )

    heading2(doc, "5.2 Tricentis Testim (industry tool named in the Dassault note)")
    body(
        doc,
        "Testim is a paid product. It stores extra information about each control so that a test can still find the button if one attribute changes. It also markets natural-language authoring. "
        "We cannot use Testim as our implementation because it is closed and licence based. It is useful as a reference: industry already accepts that plain scripts are costly to keep. Our work is a smaller, open pipeline with explicit schemas.",
    )

    heading2(doc, "5.3 Nass, Alégroth and Feldt (2024) on LLM-based localisation")
    body(
        doc,
        "This paper looks at the problem of finding the right web element when the old locator is stale. The authors show that a large language model can suggest a better match from the current page. "
        "For us this supports Objective O6. We have not plugged an LLM into localisation yet. The paper is kept as the main reference for Phase-3 healing, so that we do not invent a method that is already published without citing it.",
    )

    heading2(doc, "5.4 Generative AI for self-healing Selenium tests (ICECIT, 2025)")
    body(
        doc,
        "The ICECIT 2025 study compares open-source models with simple rules for fixing broken Selenium locators. Repair success is better with some models, but there is extra time and resource cost. "
        "This is a warning for our Gantt: if we add healing, we must measure how long a run takes, not only whether it passes. For CA-2 we only cite the result; we have not repeated their experiment.",
    )

    heading2(doc, "5.5 Reinforcement learning with dynamic XPath in Playwright (JAIGS, 2025)")
    body(
        doc,
        "A different school of work uses reinforcement learning (PPO) to pick XPath when the page changes. This is heavier than an LLM prompt and needs a training loop. "
        "We mention it to show that healing is not only an LLM problem. For a four-member B.Tech project, heuristic fallback plus a later LLM call is more realistic than training PPO in Semester VII.",
    )

    heading2(doc, "5.6 Additional baseline: Robula+ (Leotta et al.)")
    body(
        doc,
        "Before AI tools, Robula+ tried to build locators that survive small DOM edits. We treat it as the pre-AI baseline. If our parser writes a selector, we should prefer stable attributes (id, data-test) rather than long absolute XPath. That lesson is already used in the Sauce Demo mapping table in our parser.",
    )

    # 6
    heading(doc, "6. Gap in the Research / Technology / Methodology")
    body(
        doc,
        "From the review above, three lines of work are common: (i) generate Playwright/Selenium code from a plan, (ii) heal locators after a failure, and (iii) sell a commercial NL testing product. "
        "What is less common in student-accessible form is a single Quality Engineering workflow that (a) accepts a structured text case, (b) executes it immediately, (c) stores a step-level pass/fail file with a screenshot, and (d) opens a failure ticket for a named team. That combination is the gap we are targeting.",
    )
    body(
        doc,
        "A second gap is methodological. Many demos hide the schema. Mentors cannot see why a step was accepted or rejected. We locked three contracts in this month: the step schema, the report schema, and the structured prompt schema. Later LLM code must emit the same JSON. This keeps the project honest when AI is added.",
    )

    # 7
    heading(doc, "7. Description of the Proposed Solution")
    body(
        doc,
        "The proposed system is a vertical slice. A tester fills a form (or a YAML file) with site URL, feature, test id, name, objective, expected outcome and an ordered list of steps. "
        "The validator rejects empty fields. A parser turns each English line into an action enum (goto, fill, click, select, assert_text, assert_url). Playwright opens Chromium and carries out the list. If a step fails, later steps are marked skipped, a screenshot is saved, a Markdown/JSON report is written, and the notify agent looks up the module in team_ownership.yaml.",
    )
    body(
        doc,
        "The first target application is https://www.saucedemo.com/. It is public, does not need a VPN, and is widely used in QA teaching. ClickUp / Zoho will be tried only if Dassault mentors give access. Until then, Sauce Demo is the official MVP target.",
    )
    heading2(doc, "7.1 Scope for this month (in) vs later (out)")
    add_table(
        doc,
        ["In scope (done / ongoing)", "Out of scope until later phase"],
        [
            ["Structured prompt + rule parser", "LLM parser as default path"],
            ["Playwright Chromium actions listed above", "Firefox / WebKit matrix"],
            ["File reports + Streamlit demo", "Full production dashboard and DB"],
            ["Console notify + optional Slack webhook", "Jira/ClickUp ticket API"],
            ["Two structured cases + ten text drafts", "Self-healing locators in production"],
        ],
        [3.8, 3.8],
    )

    # 8
    heading(doc, "8. Requirement Analysis")
    heading2(doc, "8.1 Functional requirements")
    add_table(
        doc,
        ["ID", "Requirement", "Covered?"],
        [
            ["FR1", "Accept structured prompt and reject incomplete prompts.", "Yes"],
            ["FR2", "Parse numbered/plain steps into JSON TestSuite.", "Yes"],
            ["FR3", "Run goto, fill, click, select, assert_text, assert_url.", "Yes"],
            ["FR4", "Stop the suite after first hard failure and skip remaining steps.", "Yes"],
            ["FR5", "Save JSON and Markdown reports under data/reports.", "Yes"],
            ["FR6", "Save a screenshot for a failed step.", "Yes"],
            ["FR7", "Map module → team and emit a failure message with ticket id.", "Yes"],
            ["FR8", "Offer CLI, FastAPI and Streamlit as three ways to run a case.", "Yes"],
            ["FR9", "Allow headless and headed browser for demo vs CI.", "Yes"],
            ["FR10", "LLM interpretation of free English beyond the verb list.", "No (Phase-2)"],
        ],
        [0.8, 5.0, 1.6],
    )
    heading2(doc, "8.2 Non-functional requirements")
    add_table(
        doc,
        ["ID", "Requirement", "Current handling"],
        [
            ["NFR1", "A single Sauce Demo login flow should finish in under one minute on a college laptop.", "Observed in local runs"],
            ["NFR2", "No paid API key required for the CA-2 demo.", "Rule parser used"],
            ["NFR3", "Reports must be readable without opening the browser again.", "Markdown + JSON"],
            ["NFR4", "Secrets (Slack webhook) stay in .env, not in Git.", "Documented"],
            ["NFR5", "Code should be modular so another student can own one package.", "src/agent, executor, reporting, notify"],
        ],
        [0.8, 4.6, 2.2],
    )

    # 9
    heading(doc, "9. Technology Stack")
    add_table(
        doc,
        ["Technology", "Where it is used", "Why this choice"],
        [
            ["Python 3.11+", "All modules", "Matches AIML lab stack; Playwright and FastAPI both sit well here."],
            ["Playwright (sync API)", "Browser execution", "Named in Dassault note; better waits than older Selenium for this MVP."],
            ["Pydantic v2", "TestSuite, Report, StructuredTestPrompt", "Rejects wrong JSON early."],
            ["FastAPI + Uvicorn", "HTTP API", "Thin API for later UI; /health for a quick check."],
            ["Streamlit", "Mentor demo UI", "Enough for CA-2; React can wait."],
            ["PyYAML", "Sample cases + team map", "QA-friendly files."],
            ["JSON Schema files", "schemas/*.json", "Shared contract for parser and future LLM."],
            ["Git / GitHub", "Version control", "ESE dissemination and team merge."],
        ],
        [2.0, 2.4, 3.2],
    )

    # 10
    heading(doc, "10. Design")
    heading2(doc, "10.1 Block diagram")
    doc.add_picture(str(FIG), width=Inches(6.3))
    caption(doc, "Figure 1. End-to-end block diagram of the MVP pipeline.")

    heading2(doc, "10.2 Module design")
    add_table(
        doc,
        ["Package", "Main files", "Responsibility"],
        [
            ["src/agent", "parser.py, structured_prompt.py", "Text/YAML → TestSuite"],
            ["src/executor", "runner.py", "Playwright lifecycle and step loop"],
            ["src/reporting", "writer.py", "JSON + Markdown persistence"],
            ["src/notify", "agent.py", "Ownership lookup and alert"],
            ["src/backend", "app.py", "REST surface"],
            ["demo/", "streamlit_app.py", "Form-based demo"],
            ["config/", "team_ownership.yaml", "Module to team map"],
        ],
        [1.5, 2.4, 3.6],
    )

    heading2(doc, "10.3 Activity flow (one test run)")
    body(
        doc,
        "Start → load prompt → validate required fields → parse steps → launch Chromium → for each step: act or assert → on fail: screenshot, mark rest skipped, break → close browser → write report → if suite failed: notify agent → end. "
        "This is a linear activity. There is no parallel browser in MVP. State of a step is one of passed, failed, error, skipped.",
    )

    heading2(doc, "10.4 Data / contract design (instead of a heavy ER diagram)")
    body(
        doc,
        "MVP stores files, not a relational database. The logical records are TestSuite, Step, TestReport, StepResult and NotifyInfo. "
        "A future SQLite table can follow the same fields (run_id as key, steps as JSON). An ER diagram will be added when the database phase starts. For now the schema files are the design source.",
        first_line_indent=False,
    )
    add_table(
        doc,
        ["Record", "Key fields"],
        [
            ["Step", "id, action, description, selector, value, expected"],
            ["TestSuite", "suite_id, name, module, site_url, steps[]"],
            ["StepResult", "step_id, status, expected, actual, error, screenshot_path, duration_ms"],
            ["TestReport", "run_id, status, summary counts, steps[], notify"],
            ["NotifyInfo", "triggered, team, channel, ticket_id"],
        ],
        [1.6, 5.8],
    )

    # 11
    heading(doc, "11. Development / Implementation")
    body(
        doc,
        "Implementation in August followed a vertical-slice rule: do not wait for a perfect LLM. A regex parser was written first so that the executor and report modules could be tested with real pages. "
        "Selector guessing uses a small dictionary for Sauce Demo fields such as #user-name and #login-button. Unknown labels fall back to Playwright text= locators. This is enough for the sample cases and is documented as a limitation.",
    )
    heading2(doc, "11.1 Parser (algorithm in brief)")
    body(
        doc,
        "For each non-empty line: if the line matches Fill X with Y, emit fill; if it matches Click X, emit click; Open/Go to → goto; Verify text … is visible → assert_text; Verify URL contains → assert_url; Select A from B → select. "
        "Lines that do not match raise a validation error instead of being silently ignored. That was a deliberate choice after the mentor asked for a structured prompt.",
    )
    heading2(doc, "11.2 Executor")
    body(
        doc,
        "PlaywrightExecutor.run() creates a run_id, opens Chromium, and loops suite.steps. Assertions use expect-style checks with a timeout from settings. "
        "On failure the page screenshot is stored under data/screenshots/<run_id>/. The method returns a TestReport object; it does not print a human report by itself. That split keeps unit tests smaller.",
    )
    heading2(doc, "11.3 Reporting and notify")
    body(
        doc,
        "save_json_report and save_markdown_report write sibling files named by run_id. The Markdown file has a summary table and a step table. NotifyAgent.maybe_notify() returns the same report with notify filled. "
        "Default channel is console so that a classroom demo does not depend on Slack. If SLACK_WEBHOOK_URL is set, httpx posts the same text.",
    )
    heading2(doc, "11.4 How to run (implementation evidence)")
    body(
        doc,
        "python scripts/run_suite.py --structured tests/samples/structured/TC01_login_success.yaml runs the passing login. "
        "The same command with TC10_intentional_fail.yaml is used to show a red report and a TKT-xxxxxxxx id. Streamlit is started with streamlit run demo/streamlit_app.py. These commands are in the repository README.",
        first_line_indent=False,
    )

    # 12
    heading(doc, "12. Testing and Debugging")
    body(
        doc,
        "Testing in this period had two layers: unit tests for the parser/prompt objects, and scenario tests against Sauce Demo. Debugging was mostly locator mistakes and missing waits. Playwright’s default timeout removed the need for sleep() in the happy path.",
    )
    heading2(doc, "12.1 Test suite designed so far")
    add_table(
        doc,
        ["Case", "Type", "Expected", "Result (local)"],
        [
            ["TC01 login success", "Structured YAML", "PASS; URL has inventory.html; text Products visible", "PASS"],
            ["TC10 intentional fail", "Structured YAML", "FAIL on missing text; notify triggered", "FAIL as designed"],
            ["Parser unit tests", "pytest", "Known lines map to correct action enums", "Passing in repo"],
            ["Structured prompt validation", "pytest", "Missing field rejected", "Passing in repo"],
            ["Health API", "Manual", "GET /health returns ok", "Checked locally"],
        ],
        [1.8, 1.6, 2.6, 1.6],
    )
    heading2(doc, "12.2 Defects found and fixed")
    add_table(
        doc,
        ["Issue", "Fix"],
        [
            ["Click Login failed when the step said “Login button” vs “Login”", "Normalised aliases in selector map"],
            ["Open URL skipped if YAML already had site_url", "Executor now opens site_url before steps if no goto"],
            ["Notify fired on PASS", "maybe_notify returns early when status is passed"],
            ["Markdown table broke when error text had a pipe character", "Pipes escaped in writer.py"],
        ],
        [3.5, 4.0],
    )

    # 13
    heading(doc, "13. Project Plan and Progress")
    body(
        doc,
        "The semester plan is fifteen weeks. August mainly covers Weeks 3–8 of that plan (research freeze, executor, reports, dashboard start). "
        "Compared with the original Gantt, the vertical slice (parse–execute–report–notify) is slightly ahead. LLM parsing and self-heal are still behind, which matches the written plan and should not be shown as complete in CA-2.",
    )
    add_table(
        doc,
        ["Week (logbook)", "Planned", "Actual by 28 Aug 2026"],
        [
            ["W1–2 (13–25 Jul)", "Kickoff, roles, mentor slides", "Done: DS 1 understood; stack frozen"],
            ["W3 (27 Jul–1 Aug)", "Lit review (5), architecture, schemas", "Done"],
            ["W4 (3–8 Aug)", "Playwright bootstrap", "Done"],
            ["W5–6 (10–22 Aug)", "Parser + 3 flows MVP", "Done on Sauce Demo (login + fail demo)"],
            ["W7–8 (24 Aug–5 Sep)", "Reports + dashboard", "Reports done; Streamlit MVP done; richer dashboard still thin"],
            ["W9", "Notify agent complete", "Core done; digest mail not done"],
            ["W10–11", "Batch runner, metrics, heal start", "Not started"],
            ["W12–15", "Polish, report, poster, paper", "Not started"],
        ],
        [1.8, 2.4, 3.4],
    )
    body(
        doc,
        "Gantt (text form). Bars use # for work already done and · for remaining.",
        first_line_indent=False,
    )
    add_table(
        doc,
        ["Work item", "Jul", "Aug", "Sep", "Oct"],
        [
            ["Problem + literature", "####", "##··", "····", "····"],
            ["Schemas + architecture", "·###", "####", "····", "····"],
            ["Executor + parser MVP", "····", "####", "##··", "····"],
            ["Report + notify + demo UI", "····", "###·", "####", "····"],
            ["LLM / heal / metrics", "····", "····", "####", "##··"],
            ["Black book / poster / paper", "····", "·#··", "##··", "####"],
        ],
        [2.2, 1.2, 1.2, 1.2, 1.2],
    )

    # 14
    heading(doc, "14. Project Outcome")
    body(
        doc,
        "The B.Tech Project SOP asks every group to name the intended outcome at the start. For this industry problem the main outcome is a working prototype built under an industry statement (Dassault DS 1), with GitHub as the dissemination path and a paper draft with 3DS mentors if they agree. Deployment to a real customer tool is not promised in Semester VII.",
    )
    add_table(
        doc,
        ["Planned outcome (SOP list)", "Plan to achieve it", "Progress so far"],
        [
            [
                "Industry problem, work reviewed by industry / faculty mentors",
                "Keep DS 1 scope; demo to faculty now; demo to Dassault when the meeting is scheduled",
                "Faculty guide meetings started; Dassault meeting awaited",
            ],
            [
                "Working prototype used in a demo setting",
                "Sauce Demo + Streamlit/CLI; later optional ClickUp if access is given",
                "MVP pipeline runs TC01 and TC10",
            ],
            [
                "GitHub code with README",
                "Public repo, weekly commits",
                "Repository structured; README commands documented",
            ],
            [
                "Conference / journal draft (stretch)",
                "Use literature gap + evaluation metrics in Phase-4",
                "Five sources collected; no paper text yet",
            ],
        ],
        [2.2, 2.6, 2.8],
    )

    # 15
    heading(doc, "15. Conclusion")
    body(
        doc,
        "CA-2 records the first working cut of DS 1. The team can take a structured test case, run it on a browser, and show a pass or fail file. A failed case produces a ticket-style message for a mapped team. "
        "This is enough to discuss with the guide and, later, with Dassault. It is not the full product. The parser is still pattern based. Locator recovery is not implemented. The dashboard is only a Streamlit page.",
    )
    body(
        doc,
        "Next month (after CA-2) the group will expand sample cases, start the LLM parser behind the same JSON schema, and thicken the report view. Weekly logbook and guide meetings will continue. "
        "Similarity and AI reports for this document will be generated by the Guide and filed as Appendix A and Appendix B before the final hard copy is signed.",
    )

    # 16
    heading(doc, "16. References (APA)")
    refs = [
        "Leotta, M., Clerissi, D., Ricca, F., & Tonella, P. (2016). Approaches and tools for automated testing of web applications. In several locator-generation studies including Robula+. Journal of Software: Evolution and Process. (Robula+ locator baseline as cited in later UI-testing work.)",
        "Microsoft. (2025). Playwright Test Agents (planner, generator, healer). Playwright documentation. https://playwright.dev/docs/test-agents",
        "Nass, M., Alégroth, E., & Feldt, R. (2024). Improving web element localization by using a large language model. Software Testing, Verification and Reliability, 34(7). https://doi.org/10.1002/stvr.1893",
        "Tricentis. (n.d.). Testim: AI-powered UI test automation. https://www.testim.io/",
        "IEEE ICECIT. (2025). Generative AI for self-healing Selenium tests: Comparative evaluation of open-source LLMs and rule-based methods. https://doi.org/10.1109/icecit67774.2025.11451150",
        "Journal of Artificial Intelligence General Science. (2025). Self-healing automation with reinforcement learning: Adaptive test scripts using PPO and dynamic XPath in Playwright. https://doi.org/10.60087/jaigs.v5i1.341",
        "Dassault Systèmes. (2026). AI/ML projects with SIT – ENOVIA 2026 (internal problem brief, DS 1 Agentic Web-App Test Executor).",
        "Symbiosis Institute of Technology. (2026). B.Tech project SOP and CA-2 notice, Department of AIML, AY 2026-27.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"{i}. {ref}")
        set_run_font(r, size=12)

    # Appendices
    heading(doc, "Appendix A – Similarity Report")
    body(
        doc,
        "As per CA-2 notice dated 17 August 2026, the Similarity Report is mandatory and will be generated by the Guide. "
        "The student group will insert the signed printout after this page. Target: similarity below 10%.",
        first_line_indent=False,
    )
    p = doc.add_paragraph()
    r = p.add_run("[Attach Turnitin / college similarity PDF here — Guide generated]")
    set_run_font(r, size=12, italic=True)

    heading(doc, "Appendix B – AI Plagiarism Report")
    body(
        doc,
        "The AI plagiarism report is also mandatory and will be generated by the Guide. "
        "This monthly report was written by the group from the project repository, meeting notes and the Dassault brief. "
        "After the Guide generates the AI report, place it behind this page. The Moodle instruction requires no AI similarity.",
        first_line_indent=False,
    )
    p = doc.add_paragraph()
    r = p.add_run("[Attach AI plagiarism PDF here — Guide generated]")
    set_run_font(r, size=12, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of CA-2 Monthly Progress Report (August 2026) —")
    set_run_font(r, size=11, italic=True)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    doc.save(DOWNLOADS)
    print("Wrote", OUT)
    print("Wrote", DOWNLOADS)


if __name__ == "__main__":
    build()
